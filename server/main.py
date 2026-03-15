from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import os
import hashlib
import secrets
from datetime import datetime, timedelta
from dotenv import load_dotenv

# Load .env first, then .env.dmz as fallback
load_dotenv()
if not os.getenv("GEMINI_API_KEY") and not os.getenv("ANTHROPIC_API_KEY") and not os.getenv("OPENAI_API_KEY"):
    load_dotenv(".env.dmz")

from supabase import create_client, Client

SUPABASE_URL = os.getenv("SUPABASE_URL", "")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY", "")

app = FastAPI(title="DMZ Agents API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)


# ─── LLM Engine (self-contained) ─────────────────────────────────────────────

def get_llm_response(system_prompt: str, user_prompt: str, history: list = None) -> str:
    if history is None:
        history = []
        
    anthropic_key = os.getenv("ANTHROPIC_API_KEY")
    openai_key = os.getenv("OPENAI_API_KEY")
    gemini_key = os.getenv("GEMINI_API_KEY")

    if anthropic_key:
        from anthropic import Anthropic
        client = Anthropic(api_key=anthropic_key)
        
        msgs = []
        for h in history:
            role = "user" if h.get("role") == "user" else "assistant"
            content = h.get("content", "")
            if content.strip():
                if not msgs and role == "assistant":
                    msgs.append({"role": "user", "content": "(Sessão iniciada)"})
                if msgs and msgs[-1]["role"] == role:
                    msgs[-1]["content"] += "\n\n" + content
                else:
                    msgs.append({"role": role, "content": content})
        
        if msgs and msgs[-1]["role"] == "user":
            msgs[-1]["content"] += "\n\n" + user_prompt
        else:
            msgs.append({"role": "user", "content": user_prompt})
        
        response = client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=12288,
            system=system_prompt,
            messages=msgs
        )
        return response.content[0].text

    elif openai_key:
        from openai import OpenAI
        client = OpenAI(api_key=openai_key)
        
        msgs = [{"role": "system", "content": system_prompt}]
        for h in history:
            role = "user" if h.get("role") == "user" else "assistant"
            content = h.get("content", "")
            if content.strip():
                msgs.append({"role": role, "content": content})
        msgs.append({"role": "user", "content": user_prompt})
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=msgs,
            max_tokens=4096
        )
        return response.choices[0].message.content

    elif gemini_key:
        from google import genai
        from google.genai import types
        client = genai.Client(api_key=gemini_key)
        
        contents = []
        for h in history:
            role = "user" if h.get("role") == "user" else "model"
            content = h.get("content", "")
            if content.strip():
                if not contents and role == "model":
                    contents.append(types.Content(role="user", parts=[types.Part.from_text(text="(Sessão iniciada)")]))
                if contents and contents[-1].role == role:
                    contents[-1].parts[0].text += "\n\n" + content
                else:
                    contents.append(types.Content(role=role, parts=[types.Part.from_text(text=content)]))
                    
        if contents and contents[-1].role == "user":
            contents[-1].parts[0].text += "\n\n" + user_prompt
        else:
            contents.append(types.Content(role="user", parts=[types.Part.from_text(text=user_prompt)]))
        
        chat_model = get_model("chat", "gemini-2.0-flash")
        response = client.models.generate_content(
            model=chat_model,
            contents=contents,
            config=types.GenerateContentConfig(
                system_instruction=system_prompt,
                max_output_tokens=16384
            ),
        )
        return response.text

    else:
        raise ValueError("Nenhuma chave de LLM configurada. Por favor, configure ANTHROPIC_API_KEY, OPENAI_API_KEY ou GEMINI_API_KEY.")


def get_agent_prompt(agent_id: str) -> str:
    """Fetches agent master prompt from DB. No hardcoded fallback — prompts MUST exist in admin."""
    try:
        res = supabase.table("dmz_agents_prompts") \
            .select("content") \
            .eq("agent_id", agent_id) \
            .eq("active", True) \
            .order("version", desc=True) \
            .limit(1) \
            .execute()
        if res.data and res.data[0].get("content"):
            return res.data[0]["content"]
    except Exception as e:
        print(f"[WARN] Failed to load prompt for agent '{agent_id}': {e}")
    # No hardcoded fallback — return a minimal identity so the LLM doesn't hallucinate
    return f"Você é o agente '{agent_id}' do squad DMZ. Seu prompt ainda não foi configurado no painel admin. Responda de forma prestativa até que seu prompt seja definido."


def get_sys_prompt(p_id: str) -> str:
    """Fetches a system/tool prompt from DB. No hardcoded fallback."""
    try:
        res = supabase.table("dmz_agents_prompts") \
            .select("content") \
            .eq("agent_id", p_id) \
            .eq("active", True) \
            .limit(1) \
            .execute()
        if res.data and res.data[0].get("content"):
            return res.data[0]["content"]
    except Exception as e:
        print(f"[WARN] Failed to load sys prompt '{p_id}': {e}")
    return ""


def get_model(purpose: str, default: str = "gemini-2.0-flash") -> str:
    """Fetches the model ID for a given purpose from admin_system_models table."""
    try:
        res = supabase.table("admin_system_models") \
            .select("model_id") \
            .eq("purpose", purpose) \
            .eq("active", True) \
            .limit(1) \
            .execute()
        if res.data and res.data[0].get("model_id"):
            return res.data[0]["model_id"]
    except Exception as e:
        print(f"[WARN] Failed to load model for '{purpose}': {e}")
    return default


def process_artifacts(response_text: str) -> str:
    """Parses <dmz_artifact type="document">, generates binary docx, uploads, and injects url."""
    import re
    import tempfile
    import uuid
    import os

    # Look for document artifacts
    pattern = r'<dmz_artifact\s+type="document"\s+filename="([^"]+)"\s+title="([^"]+)">([\s\S]*?)<\/dmz_artifact>'
    
    def replacer(match):
        filename = match.group(1)
        title = match.group(2)
        content = match.group(3).strip()
        
        try:
            from docx import Document
            doc = Document()
            for line in content.split('\n'):
                l = line.strip()
                if l.startswith('# '): doc.add_heading(l[2:], level=1)
                elif l.startswith('## '): doc.add_heading(l[3:], level=2)
                elif l.startswith('### '): doc.add_heading(l[4:], level=3)
                elif l.startswith('- ') or l.startswith('* '): doc.add_paragraph(l[2:], style='List Bullet')
                elif l: doc.add_paragraph(l)

            # Save temporarily
            fd, tmp_path = tempfile.mkstemp(suffix=".docx")
            os.close(fd)
            doc.save(tmp_path)
            
            # Upload to Supabase dmz-chat-files
            bucket_name = "dmz-chat-files"
            file_key = f"artifacts/{uuid.uuid4().hex}_{filename}"
            with open(tmp_path, "rb") as f:
                supabase.storage.from_(bucket_name).upload(file_key, f)
            os.remove(tmp_path)
            
            # Retrieve public URL
            res_url = supabase.storage.from_(bucket_name).get_public_url(file_key)
            # Reconstruct tag substituting content with URL and changing type to docx_url
            return f'<dmz_artifact type="document_url" filename="{filename}" title="{title}" url="{res_url}"></dmz_artifact>'
        except Exception as e:
            print(f"[WARN] Failed to process document artifact: {e}")
            return match.group(0) # fallback to original

    return re.sub(pattern, replacer, response_text)


# ─── Routes ──────────────────────────────────────────────────────────────────

@app.get("/")
async def root():
    return {"message": "DMZ Agents API is running", "status": "healthy", "version": "2.0"}


@app.get("/health")
async def health():
    llm_provider = "none"
    if os.getenv("ANTHROPIC_API_KEY"): llm_provider = "anthropic"
    elif os.getenv("OPENAI_API_KEY"): llm_provider = "openai"
    elif os.getenv("GEMINI_API_KEY"): llm_provider = "gemini"
    return {"status": "healthy", "llm_provider": llm_provider}


@app.get("/agents")
async def get_agents():
    try:
        response = supabase.table("dmz_agents_definitions").select("*").execute()
        return response.data
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/projects")
async def get_projects():
    try:
        response = supabase.table("dmz_agents_projects").select("*").execute()
        return response.data
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class ChatRequest(BaseModel):
    message: str
    agent_id: str = "orch"
    project_id: str = "default"
    tool: str | None = None
    file_url: str | None = None
    file_type: str | None = None
    file_name: str | None = None
    files: list = []
    history: list = []


@app.post("/chat")
async def chat_interaction(req: ChatRequest):
    try:
        # Original assigned agent
        agent_db_id = req.agent_id
        if agent_db_id == "orch":
            agent_db_id = "orchestrator"

        # Build full message early for routing 
        full_message = req.message or ""

        # --- ROUTING LOGIC ---
        # If talking to orch, check if message relates to a specific specialty
        if agent_db_id == "orchestrator" and full_message.strip():
            route_prompt = f"""Analise a mensagem do usuário e decida se ela deve ser respondida pelo Orchestrator (geral) ou por um especialista do Squad DMZ.
A mensagem: '{full_message}'

Responda APENAS com o handle do agente se o tema for da alçada dele, ou "NOPE" para manter com o Orchestrator.
Especialistas e seus Domínios:
- aurora (Design: geração de imagens, logos, UI/UX visual, cores, branding, "gere uma imagem", "crie um logo", "estilo visual")
- victoria (UX: fluxo de usuário, wireframes, arquitetura de informação, usabilidade)
- ryan (Developer: backend, infra, código, APIs, bugs técnicos, "escreva um script", "como fazer tal função")
- theron (Legal: contratos, análise jurídica, LGPD, termos de uso)
- cassandra (Copy: textos persuasivos, marketing, legendas, roteiros, emails frios)
- lucas (PO: produto, estratégia, roadmap, requisitos de negócio)
- sofia (DB: bancos de dados, queries, arquitetura de dados)

Regra: Se a mensagem mencionar "gerar imagem" ou "crie uma foto/ilustração", mande obrigatoriamente para @aurora."""
            
            try:
                # Fast classification (use a smaller model if needed, but here we reuse the response engine)
                route_decision = get_llm_response("Você é um roteador de squad. Responda apenas com o handle ou NOPE.", route_prompt).strip().lower()
                valid_handles = ["aurora", "victoria", "ryan", "theron", "cassandra", "lucas", "sofia"]
                
                # Force aurora for image-like requests if the LLM is being lazy
                if not any(h in route_decision for h in valid_handles):
                    img_triggers = ["imagem", "foto", "ilustra", "logo", "desenho", "image", "picture", "create image"]
                    if any(t in full_message.lower() for t in img_triggers):
                        route_decision = "aurora"

                if any(h in route_decision for h in valid_handles):
                    for h in valid_handles:
                        if h in route_decision:
                            # Map handle to DB ID
                            mapping = {
                                "aurora": "design_chief",
                                "victoria": "ux",
                                "ryan": "developer",
                                "theron": "legal_chief",
                                "cassandra": "copy_chief",
                                "lucas": "po",
                                "sofia": "db_sage"
                            }
                            agent_db_id = mapping.get(h, agent_db_id)
                            # Update parameters so the frontend sees the change
                            req.agent_id = h
                            # Add routing context for the specialist
                            full_message = f"[CONTEXTO: Especialista {h} chamado. Execute a tarefa agora de forma direta. Se houver comandos de ferramentas (como gerar imagem ou código), execute-os imediatamente.]\n\n{full_message}"
                            # Mantém todo o histórico intacto para contexto linear
                            req.history = req.history
                            print(f"[routing] Rerouted to specialist: {h}")
                            break
            except Exception as e:
                print(f"[routing] Error: {e}")
                pass

        # Build system prompt from DB
        system_prompt = get_agent_prompt(agent_db_id)

        # Apply global formatting rule
        formatting = get_sys_prompt("system_formatting")
        if formatting:
            system_prompt += "\n" + formatting

        # Inject current date/time in GMT-3 (São Paulo)
        from datetime import datetime, timezone, timedelta
        tz_sp = timezone(timedelta(hours=-3))
        now_sp = datetime.now(tz_sp)
        day_names = ["segunda-feira", "terça-feira", "quarta-feira", "quinta-feira", "sexta-feira", "sábado", "domingo"]
        month_names = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
        day_name = day_names[now_sp.weekday()]
        month_name = month_names[now_sp.month - 1]
        date_str = f"{day_name}, {now_sp.day} de {month_name} de {now_sp.year}"
        time_str = now_sp.strftime("%H:%M")
        
        system_prompt += (
            f"\n\nCONTEXTO TEMPORAL E CONTINUIDADE:"
            f"\n- Data atual: {date_str}"
            f"\n- Horário atual: {time_str} (fuso horário: GMT-3, América/São Paulo)"
            f"\n- CONTINUIDADE DO SQUAD: Você está atuando em uma linha do tempo cooperativa com outros especialistas no chat. Se houver mensagens anteriores no histórico começando com algo como [Aurora]: ou [Lucas]:, saiba que essa mensagem não foi dita pelo usuário, mas sim pelo seu COB (Colega de Squad). Leia TUDO o que foi falado antes e continue o projeto a partir do último ponto de forma harmoniosa."
        )

        # CRITICAL: Behavior rules
        system_prompt += (
            "\n\nREGRAS ABSOLUTAS (NUNCA VIOLAR):"
            "\n- NUNCA mencione que foi 'acionado', 'chamado' ou que é uma ferramenta. Aja como o profissional especializado."
            "\n- Execute as tarefas imediatamente. Se pedirem imagem, use comandos de imagem (implícitos no sistema)."
            "\n- Ao citar colegas, use @nome (ex: @aurora, @lucas)."
            "\n- Responda de forma executiva, direta e focada em resultados de alto nível."
        )

        # Artifacts Rendering Rules
        system_prompt += (
            "\n\nGERAÇÃO DE ARTEFATOS E ARQUIVOS (MANDATÓRIO):"
            "\nSe a sua resposta envolver criar um arquivo completo (Landing Page, Documento, Componente UI, etc), VOCÊ DEVE encapsular esse código integralmente dentro de tags XML `<dmz_artifact>` na sua resposta. NUNCA envie HTML quebrado no chat. Formato exato:"
            "\n<dmz_artifact type=\"(html|jsx|document)\" filename=\"nome-do-arquivo.ext\" title=\"Título Legível\">"
            "\n   [Todo o código aqui]"
            "\n</dmz_artifact>"
            "\n"
            "\nIMPORTANTE: NUNCA envolva a tag <dmz_artifact> em blocos de código markdown (```). Use as tags XML diretamente no seu texto."
            "\n"
            "\n## REGRAS ESTÉTICAS ABSOLUTAS PARA HTML/UI (NÍVEL PREMIUM/LUXURY):"
            "\nO usuário do DMZ abomina UIs genéricas (ex: tela branca com Inter e botões azuis simples). Você deve surpreendê-lo com design digno de prêmios (Awwwards)."
            "\n"
            "\n1. Tipografia de Alto Padrão: Sempre importe do Google Fonts fontes de impacto. Use fontes Display para títulos (Ex: 'Montserrat', 'Playfair Display', 'Clash Display', 'Outfit') (weight: 600-800) e fontes Sans limpas para textos (weight: 300-400). Jamais use Roboto, Arial ou system-ui genérica."
            "\n2. Sistema de Cores e Fundo (Luxury Dark Theming encorajado): NUNCA use #000000 ou #ffffff puro no background principal. Use cinzas dinâmicos/matizes suaves (#0A0A0F, #FAFAF7, etc). Use no máximo 1 cor vibrante como Acento (Accent) com variações de opacidade (rgba)."
            "\n3. Efeitos Atmosféricos em Dark Themes (OBRIGATÓRIO usar Tailwind ou CSS In-line):"
            "\n   - Se criar fundos escuros, obrigatoriamente gere uma textura de \"noise\" suave (pointer-events: none, opacity: 0.1) cobrindo a tela."
            "\n   - Construa um brilho ambiente (Glow) invisível rodando ao fundo usando pseudo elementos com \"radial-gradient\" sutis nas pontas (#accent rgba)."
            "\n4. Bordas Elegantes: Jamais use bordas cinzas sólidas e grossas em dark mode. Use SEMPRE rgba branco ultra suave [border: 1px solid rgba(255,255,255,0.08)]."
            "\n5. Animações Modernas e Orgânicas (CÓDIGO OBRIGATÓRIO PARA LANDING PAGES):"
            "\n   - Adicione no corpo da sua página JS um IntersectionObserver. As seções (hero, blocos de cards, textos) devem brotar suavemente no scroll (opacity 0->1, translateY 20-30px) de forma escalonada (delay)."
            "\n   - Hover obrigatório em elementos interativos com transition suave (0.3s ease-out), como glows nos cards via border-color ou shadow."
            "\n   - Só anime \"opacity\" e \"transform\". NUNCA use loops infinitos nervosos nem anime \"width/margin/left\"."
            "\n6. Arquitetura Independente: Adicione Tailwind via script CDN (<script src=\"https://cdn.tailwindcss.com\"></script>), inclua os ícones SVG via lucide-react CDN."
            "\n7. Tudo precisa funcionar auto-contido. O documento HTML que você devolver dentro de <html> até </html> não pode depender de mais nada local."
            "\n"
            "\n## ESTRUTURA E COMPLEXIDADE (LANDING PAGES COMPLETAS):"
            "\n- NUNCA crie apenas uma seção 'Hero'. Se o usuário pediu uma landing page, ele espera um fluxo de conversão completo."
            "\n- OBRIGATÓRIO ter no mínimo 5 seções distintas: 1. Hero, 2. Problema/Solução ou Features detalhadas, 3. Social Proof/Depoimentos ou Números, 4. FAQ ou Processo, 5. CTA Final e Footer."
            "\n- Use espaçamentos generosos (py-24, py-32) entre seções. Dê respiro ao design."
            "\n- Cada seção deve ter um layout visualmente diferente da anterior (ex: zig-zag de imagens/texto, grids de cards, seções full-width com background accent)."
            "\n- NÃO SEJA PREGUIÇOSO. Escreva o código completo para uma experiência de navegação rica."
            "\n"
            "\n## REGRAS PARA DOCUMENTOS (PDF, WORD, CONTRATOS, PROPOSTAS):"
            "\n- NUNCA diga que não consegue criar arquivos PDF ou Word diretamente. Se o usuário pedir um PDF ou Word, você deve gerar uma versão HTML/CSS altamente formatada como um documento profissional e entregar via `<dmz_artifact type=\"html\" ...>`. O usuário pode imprimir ou salvar o que você renderizar."
            "\n- Design de Documento: Use fundo branco (#FFFFFF), margens generosas (padding: 40px), tipografia serifada profissional (Ex: 'Lora', 'Playfair Display') para títulos e sans-serif limpa para o corpo."
            "\n- Estrutura: Inclua sempre cabeçalho com logo/título, data, seções numeradas e rodapé."
        )


        # Handle tool logic
        # Auto-detect image generation from natural language (if no tool explicitly selected)
        effective_tool = req.tool
        if not effective_tool and full_message.strip():
            import re
            msg_lower = full_message.lower()
            # Remove internal transcription marker for detection
            clean_for_detection = msg_lower.replace("[transcrição interna do áudio]:", "").strip()
            # Remove routing context
            if "[contexto:" in clean_for_detection:
                clean_for_detection = clean_for_detection.split("]\n\n", 1)[-1] if "]\n\n" in clean_for_detection else clean_for_detection
            
            # Regex patterns for image generation requests
            img_patterns = [
                r'ger[ea]\s+\d*\s*imagen[s]?',       # gere/gera 3 imagens, gere imagens
                r'cri[ea]\s+\d*\s*imagen[s]?',        # crie/cria 3 imagens
                r'fa[çz][ao]\s+\d*\s*imagen[s]?',     # faça 3 imagens
                r'ger[ea]r?\s+\d*\s*imagen[s]?',      # gerar imagens
                r'ger[ea]r?\s+\d*\s*foto[s]?',        # gerar fotos
                r'cri[ea]r?\s+\d*\s*foto[s]?',        # criar fotos
                r'ger[ea]r?\s+(uma?\s+)?imagem',       # gerar uma imagem
                r'cri[ea]r?\s+(uma?\s+)?imagem',       # criar uma imagem
                r'fa[çz][ao]\s+(uma?\s+)?imagem',      # faça uma imagem
                r'ger[ea]r?\s+(um\s+)?logo',           # gerar logo
                r'cri[ea]r?\s+(uma?\s+)?arte',         # criar arte
                r'desenh[ea]r?\s',                     # desenhe/desenhar
                r'generate\s+\d*\s*image',             # generate image
                r'create\s+\d*\s*image',               # create image
                r'make\s+\d*\s*image',                 # make image
                r'draw\s',                             # draw
                r'ilustr[ea]r?\s',                     # ilustrar/ilustre
                r'ger[ea]r?\s+\d*\s*ilustra',          # gerar ilustração
            ]
            if any(re.search(p, clean_for_detection) for p in img_patterns):
                effective_tool = "createImage"
                print(f"[chat] Auto-detected image generation request: {clean_for_detection[:60]}")

        if effective_tool == "searchWeb":
            tool_prompt = get_sys_prompt("tool_search_web")
            if tool_prompt:
                system_prompt += "\n" + tool_prompt
            firecrawl_key = os.getenv("FIRECRAWL_API_KEY")
            if firecrawl_key:
                try:
                    import requests as req_lib
                    search_query_prompt = f"Usuário ordenou pesquisa web sobre: '{full_message}'. Escreva APENAS o termo de busca (sem citações)."
                    search_query = get_llm_response("Seja direto.", search_query_prompt).strip()
                    fc_resp = req_lib.post(
                        "https://api.firecrawl.dev/v1/search", 
                        json={"query": search_query, "limit": 4}, 
                        headers={"Authorization": f"Bearer {firecrawl_key}", "Content-Type": "application/json"},
                        timeout=15
                    )
                    if fc_resp.status_code == 200:
                        results = fc_resp.json().get("data", [])
                        snippets = "\n\n".join([f"[{r.get('title')}]({r.get('url')}): {r.get('description', '')}" for r in results])
                        full_message += f"\n\n[CONTEXTO RETORNADO DA WEB VIA FIRECRAWL PARA '{search_query}']:\n{snippets[:3000]}"
                    else:
                        full_message += f"\n\n[Firecrawl API falhou com status {fc_resp.status_code}]"
                except Exception as e:
                    full_message += f"\n\n[Erro ao tentar Firecrawl: {str(e)}]"
        elif effective_tool == "createImage":
            gemini_key = os.getenv("GEMINI_API_KEY")
            if gemini_key:
                from google import genai
                from google.genai import types
                import uuid
                import re as re_mod
                import base64
                
                # Clean internal markers from the message for cleaner prompt
                clean_msg = full_message.replace("[Transcrição Interna do Áudio]:", "").strip()
                if "[CONTEXTO:" in clean_msg:
                    clean_msg = clean_msg.split("]\n\n", 1)[-1] if "]\n\n" in clean_msg else clean_msg
                
                # Detect how many images requested (default 1, max 4)
                num_match = re_mod.search(r'(\d+)\s*imagen[s]?', clean_msg.lower())
                if not num_match:
                    num_match = re_mod.search(r'(\d+)\s*foto[s]?', clean_msg.lower())
                if not num_match:
                    num_match = re_mod.search(r'(\d+)\s*ilustra', clean_msg.lower())
                num_images = min(int(num_match.group(1)), 4) if num_match else 1
                
                # Get model from admin (default: gemini-3-pro-image-preview)
                image_model = get_model("image_generation", "gemini-3-pro-image-preview")
                print(f"[chat] Using image model: {image_model}, generating {num_images} image(s)")
                
                try:
                    client = genai.Client(api_key=gemini_key)
                    generated_urls = []
                    
                    for i in range(num_images):
                        # Use generate_content API (required for gemini-3-pro-image-preview)
                        img_prompt = f"Generate a high quality, photorealistic image: {clean_msg}"
                        if num_images > 1:
                            img_prompt += f" (variation {i+1} of {num_images}, make each unique)"
                        
                        response = client.models.generate_content(
                            model=image_model,
                            contents=img_prompt,
                            config=types.GenerateContentConfig(
                                response_modalities=['TEXT', 'IMAGE'],
                            ),
                        )
                        
                        # Extract image from response parts
                        if response.candidates and response.candidates[0].content and response.candidates[0].content.parts:
                            for part in response.candidates[0].content.parts:
                                if part.inline_data and part.inline_data.mime_type and part.inline_data.mime_type.startswith("image/"):
                                    img_bytes = part.inline_data.data
                                    ext = "png" if "png" in part.inline_data.mime_type else "jpg"
                                    file_name = f"generated/{uuid.uuid4()}.{ext}"
                                    content_type = part.inline_data.mime_type
                                    supabase.storage.from_("images").upload(file_name, img_bytes, {"content-type": content_type})
                                    img_url = supabase.storage.from_("images").get_public_url(file_name)
                                    generated_urls.append(img_url)
                                    break  # one image per iteration
                    
                    if not generated_urls:
                        raise ValueError("O modelo não retornou imagens na resposta.")
                    
                    images_md = "\n".join([f"![Imagem {i+1}]({url})" for i, url in enumerate(generated_urls)])
                    count_word = f"{len(generated_urls)} imagens" if len(generated_urls) > 1 else "a imagem"
                    msg_surround = get_llm_response(
                        system_prompt,
                        f"Você acabou de gerar {count_word} com sucesso para o usuário. "
                        f"Dê uma resposta muito breve e amigável apresentando. "
                        f"NUNCA mencione nomes técnicos de ferramentas ou IDs internos. O pedido original foi: {clean_msg}"
                    )
                    return {
                        "agent_id": req.agent_id,
                        "content": f"{msg_surround}\n\n{images_md}",
                        "status": "success"
                    }
                except Exception as e:
                    print(f"[chat] Image generation error: {e}")
                    return {
                        "agent_id": req.agent_id,
                        "content": f"Não foi possível gerar a imagem neste momento. Erro: {str(e)}",
                        "status": "error"
                    }
            else:
                system_prompt += "\nMODO GERAÇÃO DE IMAGEM: A API do Google (Gemini API_KEY) não está configurada, então apenas descreva a imagem."
        elif effective_tool == "writeCode":
            code_prompt = get_sys_prompt("tool_write_code")
            if code_prompt:
                system_prompt += "\n" + code_prompt

        # Process all files
        all_files = list(req.files) if req.files else []
        if req.file_url:
            all_files.append({"url": req.file_url, "type": req.file_type or "", "name": req.file_name or ""})

        has_image = any("image" in f.get("type", "").lower() or (f.get("url", "").split(".")[-1].lower() in ["jpg", "jpeg", "png", "gif", "webp", "heic", "avif"]) for f in all_files)
        has_audio = any("audio" in f.get("type", "").lower() or (f.get("url", "").split(".")[-1].lower() in ["mp3", "wav", "ogg", "m4a", "aac"]) for f in all_files)
        has_pdf = any("pdf" in f.get("type", "").lower() or (f.get("url", "").split(".")[-1].lower() == "pdf") for f in all_files)

        if has_image:
            img_prompt = get_sys_prompt("attachment_image")
            if img_prompt: system_prompt += "\n" + img_prompt
        if has_audio:
            audio_prompt = get_sys_prompt("attachment_audio")
            if audio_prompt: system_prompt += "\n" + audio_prompt
        if has_pdf:
            pdf_prompt = get_sys_prompt("attachment_pdf")
            if pdf_prompt: system_prompt += "\n" + pdf_prompt

        image_files = []
        file_contexts = []
        for f in all_files:
            ftype = f.get("type", "").lower()
            fext = f.get("url", "").split(".")[-1].lower() if f.get("url") else ""
            fname = f.get("name", "")
            if "image" in ftype or fext in ["jpg", "jpeg", "png", "gif", "webp", "heic", "avif"]:
                image_files.append(f)
            elif "audio" in ftype or fext in ["mp3", "wav", "ogg", "m4a", "aac"]:
                file_contexts.append(f"[Áudio recebido ({fname}): {f.get('url')}]")
            else:
                file_contexts.append(f"[Arquivo recebido ({fname}): {f.get('url')} - Tipo: {f.get('type')}]")

        if file_contexts:
            full_message += "\n\n" + "\n".join(file_contexts)

        # For multiple image attachments: use multimodal LLM
        if len(image_files) > 0:
            gemini_key = os.getenv("GEMINI_API_KEY")
            anthropic_key = os.getenv("ANTHROPIC_API_KEY")

            if gemini_key:
                from google import genai
                from google.genai import types
                import requests as req_lib
                
                client = genai.Client(api_key=gemini_key)

                img_parts = []
                for img in image_files:
                    try:
                        img_resp = req_lib.get(img.get("url"), timeout=10)
                        img_parts.append(types.Part.from_bytes(data=img_resp.content, mime_type=img.get("type") or "image/jpeg"))
                    except Exception as e:
                        print(f"Failed to fetch image {img.get('url')}: {e}")

                contents = []
                for h in req.history:
                    role_h = "user" if h.get("role") == "user" else "model"
                    content_h = h.get("content", "")
                    if content_h.strip():
                        if not contents and role_h == "model":
                            contents.append(types.Content(role="user", parts=[types.Part.from_text(text="(Sessão iniciada)")]))
                        if contents and contents[-1].role == role_h:
                            contents[-1].parts[0].text += "\n\n" + content_h
                        else:
                            contents.append(types.Content(role=role_h, parts=[types.Part.from_text(text=content_h)]))
                
                prompt_parts = img_parts.copy()
                if full_message:
                    prompt_parts.append(types.Part.from_text(text=full_message))
                
                if contents and contents[-1].role == "user":
                    contents[-1].parts.extend(prompt_parts)
                else:
                    contents.append(types.Content(role="user", parts=prompt_parts))

                try:
                    response = client.models.generate_content(
                        model='gemini-2.0-flash',
                        contents=contents,
                        config=types.GenerateContentConfig(system_instruction=system_prompt)
                    )
                    return {"agent_id": req.agent_id, "content": response.text, "status": "success"}
                except Exception as e:
                    return {"agent_id": req.agent_id, "content": f"Erro na análise das imagens com Gemini: {str(e)}", "status": "error"}

            elif anthropic_key:
                import base64, requests as req_lib
                from anthropic import Anthropic
                
                client = Anthropic(api_key=anthropic_key)
                msgs = []
                for h in req.history:
                    role_h = "user" if h.get("role") == "user" else "assistant"
                    content_h = h.get("content", "")
                    if content_h.strip():
                        if not msgs and role_h == "assistant":
                            msgs.append({"role": "user", "content": "(Sessão iniciada)"})
                        if msgs and msgs[-1]["role"] == role_h:
                            msgs[-1]["content"] += "\n\n" + content_h
                        else:
                            msgs.append({"role": role_h, "content": content_h})
                
                content_block = []
                for img in image_files:
                    try:
                        img_resp = req_lib.get(img.get("url"), timeout=10)
                        img_b64 = base64.standard_b64encode(img_resp.content).decode("utf-8")
                        content_block.append({
                            "type": "image", 
                            "source": {"type": "base64", "media_type": img.get("type") or "image/jpeg", "data": img_b64}
                        })
                    except Exception as e:
                        print(f"Failed to fetch image {img.get('url')}: {e}")
                
                if full_message:
                    content_block.append({"type": "text", "text": full_message})
                
                if msgs and msgs[-1]["role"] == "user":
                    if isinstance(msgs[-1]["content"], str):
                        msgs[-1]["content"] = [{"type": "text", "text": msgs[-1]["content"]}] + content_block
                    else:
                        msgs[-1]["content"].extend(content_block)
                else:
                    msgs.append({"role": "user", "content": content_block})

                response = client.messages.create(
                    model="claude-3-5-sonnet-20241022",
                    max_tokens=4096,
                    system=system_prompt,
                    messages=msgs
                )
                return {"agent_id": req.agent_id, "content": response.content[0].text, "status": "success"}

        # Call LLM (standard response)
        response_text = get_llm_response(system_prompt, full_message, req.history)

        # --- ARTIFACT SPECIALIST PASS (Claude for Quality) ---
        # If the response mentions creating a file, artifact or layout, we can trigger a specialist pass
        # for maximum quality in HTML/PDF/JSX using the dedicated Claude key
        if any(keyword in response_text.lower() for keyword in ["<dmz_artifact", "html", "contrato", "proposta", "site", "landing page", "código", "script"]):
            try:
                # Fetch artifact specialist config
                res_spec = supabase.table("admin_system_models").select("*").eq("purpose", "artifact_generation").eq("active", True).limit(1).execute()
                if res_spec.data:
                    spec_config = res_spec.data[0]
                    spec_key = spec_config.get("config", {}).get("api_key")
                    if spec_key:
                        from anthropic import Anthropic
                        client_spec = Anthropic(api_key=spec_key)
                        
                        spec_prompt = f"""Você é o Especialista em Engenharia de Código e Design do Squad DMZ.
Sua tarefa é RECREAR ou REFINAR o artefato técnico (HTML/CSS/DOC) contido ou sugerido na mensagem abaixo para garantir nível PREMIUM (Awwwards-quality).

Mensagem original:
---
{response_text}
---

INSTRUÇÕES DE QUALIDADE:
1. Extraia a intenção do arquivo e crie o código INTEGRAL, completo, extenso e de alta qualidade.
2. O HTML gerado roda dentro de um iframe isolado. Carregue TUDO via CDN — sem bundler, sem imports de módulo.
3. NÃO entregue esqueletos ou placeholders. Gere seções COMPLETAS com conteúdo real e visual rico.

DESIGN SYSTEM OBRIGATÓRIO (carregar via CDN no <head>):

TIPOGRAFIA (SEMPRE incluir):
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Sans:wght@300;400;500;600;700&family=JetBrains+Mono:wght@300;400&display=swap" rel="stylesheet">

CSS UTILITÁRIOS (SEMPRE incluir):
<script src="https://cdn.tailwindcss.com"></script>
<script>
  tailwind.config = {{{{
    theme: {{{{ extend: {{{{
      fontFamily: {{{{
        display: ['Syne', 'sans-serif'],
        sans: ['DM Sans', 'sans-serif'],
        mono: ['JetBrains Mono', 'monospace'],
      }}}}
    }}}}}}}}
  }}}}
</script>

ANIMAÇÕES (incluir quando houver scroll ou transições visuais):
<script src="https://cdnjs.cloudflare.com/ajax/libs/gsap/3.12.5/gsap.min.js"></script>
<script src="https://cdnjs.cloudflare.com/ajax/libs/gsap/3.12.5/ScrollTrigger.min.js"></script>
Use GSAP para scroll reveal: gsap.registerPlugin(ScrollTrigger); e aplique em elementos com classe .reveal

ÍCONES (incluir quando houver ícones):
<script src="https://unpkg.com/lucide@latest/dist/umd/lucide.js"></script>
<script>lucide.createIcons();</script>
Use: <i data-lucide="nome-do-icone"></i> — NUNCA use emojis como ícones funcionais.

INTERATIVIDADE (incluir para tabs, modais, dropdowns, toggles):
<script defer src="https://cdn.jsdelivr.net/npm/alpinejs@3.x.x/dist/cdn.min.js"></script>

REGRAS DE DESIGN:
- Use gradientes sutis, glassmorphism (backdrop-blur), sombras premium (shadow-2xl)
- Micro-animações em hover (scale, translate, glow)
- Tipografia hierárquica: font-display para títulos grandes, font-sans para corpo
- Paleta de cores harmoniosa — NUNCA use cores genéricas puras (red, blue, green)
- Espaçamento generoso (py-20, py-32) entre seções
- Mobile-first responsive design com breakpoints Tailwind (sm:, md:, lg:, xl:)
- Backgrounds texturizados ou com gradientes — NUNCA deixe fundo branco liso genérico
- Mínimo 5-7 seções completas para landing pages (hero, features, about, testimonials, CTA, footer, etc)
- Conteúdo textual realista e profissional — NUNCA use "Lorem ipsum"
- Se for um contrato/proposta, use layout de papel A4 elegante com bordas e tipografia serifada.

ESTRUTURA OBRIGATÓRIA DO HTML:
- Comece com <!DOCTYPE html> completo
- Inclua <meta charset="utf-8"> e <meta name="viewport">
- Carregue CDNs no <head> na ordem: fonts → tailwind → gsap → lucide → alpine
- Chame lucide.createIcons() e gsap no final do <body>

RETORNO:
5. Retorne APENAS o texto da mensagem original, mas substitua ou adicione as tags <dmz_artifact type="..." filename="..." title="...">CÓDIGO HTML COMPLETO E REFINADO</dmz_artifact> com o código premium.
6. Mantenha o tom da mensagem original, apenas turbine ao máximo a parte técnica e visual.
7. Gere código EXTENSO e COMPLETO — páginas profissionais têm 300-800 linhas de código.
"""
                        
                        response_spec = client_spec.messages.create(
                            model=spec_config.get("model_id", "claude-3-5-sonnet-20241022"),
                            max_tokens=16384,
                            messages=[{"role": "user", "content": spec_prompt}]
                        )
                        if response_spec.content and len(response_spec.content[0].text) > 100:
                            response_text = response_spec.content[0].text
                            print("[artifacts] Enhanced with Claude Specialist Pass")
            except Exception as e:
                print(f"[artifacts] Specialist pass failed: {e}")

        # Post-process generated artifacts to native formats (Docx, etc)
        response_text = process_artifacts(response_text)

        return {
            "agent_id": req.agent_id,
            "content": response_text,
            "status": "success"
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── Audio Transcription Endpoint ─────────────────────────────────────────────

class TranscribeRequest(BaseModel):
    audio_url: str
    file_name: str = "audio.webm"

@app.post("/transcribe")
async def transcribe_audio(req: TranscribeRequest):
    """
    Downloads audio from URL, uploads to Gemini Files API,
    and returns full transcription. Supports files up to 2 hours.
    """
    import requests as req_lib
    import tempfile
    import time

    gemini_key = os.getenv("GEMINI_API_KEY")
    if not gemini_key:
        raise HTTPException(status_code=500, detail="GEMINI_API_KEY not configured")

    try:
        # 1. Download audio from Supabase
        print(f"[transcribe] Downloading: {req.audio_url[:80]}...")
        audio_resp = req_lib.get(req.audio_url, timeout=120, stream=True)
        if audio_resp.status_code != 200:
            raise HTTPException(status_code=400, detail=f"Failed to download audio: HTTP {audio_resp.status_code}")

        # Determine MIME type
        ext = req.file_name.rsplit(".", 1)[-1].lower() if "." in req.file_name else "webm"
        mime_map = {
            "mp3": "audio/mpeg", "wav": "audio/wav", "m4a": "audio/m4a",
            "ogg": "audio/ogg", "opus": "audio/opus", "flac": "audio/flac",
            "aac": "audio/aac", "webm": "audio/webm", "3gp": "audio/3gpp",
            "amr": "audio/amr",
        }
        mime_type = mime_map.get(ext, f"audio/{ext}")

        # 2. Save to temp file
        with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as tmp:
            for chunk in audio_resp.iter_content(chunk_size=8192):
                tmp.write(chunk)
            tmp_path = tmp.name

        file_size = os.path.getsize(tmp_path)
        print(f"[transcribe] Downloaded {file_size / 1024 / 1024:.1f} MB as {mime_type}")

        # 3. Upload to Gemini Files API
        from google import genai
        from google.genai import types
        client = genai.Client(api_key=gemini_key)

        print(f"[transcribe] Uploading to Gemini Files API...")
        uploaded_file = client.files.upload(
            file=tmp_path,
            config={"mime_type": mime_type, "display_name": req.file_name}
        )

        # 4. Wait for file processing (Gemini processes async for large files)
        print(f"[transcribe] File uploaded: {uploaded_file.name}, state: {uploaded_file.state}")
        max_wait = 300  # 5 minutes max wait
        waited = 0
        while uploaded_file.state.name == "PROCESSING" and waited < max_wait:
            time.sleep(5)
            waited += 5
            uploaded_file = client.files.get(name=uploaded_file.name)
            print(f"[transcribe] Waiting... state: {uploaded_file.state} ({waited}s)")

        if uploaded_file.state.name == "FAILED":
            raise HTTPException(status_code=500, detail="Gemini failed to process the audio file")

        # 5. Transcribe using Gemini — load prompt from admin DB
        print(f"[transcribe] Requesting transcription...")
        
        # Fetch transcription prompt from DB (voice_transcription)
        transcription_prompt = get_sys_prompt("voice_transcription")
        if not transcription_prompt or len(transcription_prompt) < 50:
            # Fallback with strong paragraph formatting
            transcription_prompt = (
                "Você é o módulo de transcrição de áudio do DMZ.\n"
                "Sua função é processar e transcrever áudios enviados pelo usuário com máxima fidelidade.\n\n"
                "REGRAS:\n"
                "- Mantenha a transcrição fiel ao que foi dito, preservando pausas e ênfases quando relevante.\n"  
                "- Corrija erros gramaticais óbvios de fala, mas mantenha o tom e vocabulário original.\n"
                "- Se o áudio estiver inaudível em trechos, indique [inaudível] no ponto exato.\n"
                "- Se houver múltiplos falantes, tente identificá-los como Falante 1, Falante 2, etc.\n"
                "- Após a transcrição, ofereça um resumo dos pontos principais se o áudio for longo (>2min)."
            )
        
        # Always append strong paragraph formatting instruction
        paragraph_instruction = (
            "\n\nFORMATAÇÃO OBRIGATÓRIA:\n"
            "- Você DEVE separar a transcrição em parágrafos curtos de no máximo 3-4 frases cada.\n"
            "- Use DUAS quebras de linha (\\n\\n) entre cada parágrafo.\n"
            "- Quando o assunto mudar, inicie um novo parágrafo.\n"
            "- Quando houver pausa natural na fala, inicie um novo parágrafo.\n"
            "- NUNCA retorne a transcrição como um bloco contínuo de texto.\n"
            "- O resultado deve ser fácil de ler, com parágrafos bem definidos."
        )
        
        full_transcription_prompt = transcription_prompt + paragraph_instruction
        
        transcription_model = get_model("transcription", "gemini-2.0-flash")
        response = client.models.generate_content(
            model=transcription_model,
            contents=[
                types.Part.from_uri(
                    file_uri=uploaded_file.uri,
                    mime_type=mime_type
                ),
                full_transcription_prompt
            ],
            config=types.GenerateContentConfig(
                max_output_tokens=65536,
            )
        )

        transcription = response.text
        print(f"[transcribe] Done! Transcription length: {len(transcription)} chars")

        # 6. Cleanup
        try:
            os.unlink(tmp_path)
            client.files.delete(name=uploaded_file.name)
        except Exception:
            pass

        return {
            "transcription": transcription,
            "duration_estimate": f"{file_size / 1024 / 16:.0f}s",  # rough estimate
            "char_count": len(transcription),
            "status": "success"
        }

    except Exception as e:
        print(f"[transcribe] Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ─── API Keys ──────────────────────────────────────────────────────────────────

import hashlib
import secrets
from fastapi import Header

class GenerateKeyRequest(BaseModel):
    project_id: str

@app.post("/api/keys/generate")
async def generate_key(req: GenerateKeyRequest, authorization: str = Header(None)):
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid Authorization header")
    token = authorization.split(" ")[1]
    
    try:
        user_resp = supabase.auth.get_user(token)
        if not user_resp or not user_resp.user:
            raise HTTPException(status_code=401, detail="Invalid token")
        user_id = user_resp.user.id
    except Exception as e:
        raise HTTPException(status_code=401, detail=f"Auth error: {str(e)}")
    
    # Check ownership
    proj_resp = supabase.table("dmz_agents_projects").select("id").eq("id", req.project_id).eq("owner_id", user_id).execute()
    if not proj_resp.data:
        raise HTTPException(status_code=403, detail="Not authorized to generate key for this project")
    
    raw_key = f"dmz_pk_{secrets.token_urlsafe(32)}"
    key_hash = hashlib.sha256(raw_key.encode()).hexdigest()
    key_prefix = raw_key[:12]
    
    new_key = {
        "project_id": req.project_id,
        "owner_id": user_id,
        "key_hash": key_hash,
        "key_prefix": key_prefix,
        "name": "CLI Security Key"
    }
    
    try:
        res = supabase.table("dmz_agents_apikeys").insert(new_key).execute()
        return {"key": raw_key, "prefix": key_prefix, "inserted_at": res.data[0]["created_at"]}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save key: {str(e)}")

class ValidateKeyRequest(BaseModel):
    api_key: str
    project_slug: str

@app.post("/api/validate-key")
async def validate_key(req: ValidateKeyRequest):
    key_hash = hashlib.sha256(req.api_key.encode()).hexdigest()
    
    proj_resp = supabase.table("dmz_agents_projects").select("id").eq("slug", req.project_slug).execute()
    if not proj_resp.data:
        raise HTTPException(status_code=404, detail="Project not found")
    project_id = proj_resp.data[0]["id"]
    
    key_resp = supabase.table("dmz_agents_apikeys").select("id, is_active").eq("key_hash", key_hash).eq("project_id", project_id).execute()
    
    if not key_resp.data or not key_resp.data[0]["is_active"]:
         raise HTTPException(status_code=401, detail="Invalid or inactive API Key for this project")
         
    # Update last_used_at without failing if it breaks (best effort)
    try:
        supabase.table("dmz_agents_apikeys").update({"last_used_at": "now()"}).eq("id", key_resp.data[0]["id"]).execute()
    except Exception:
        pass
          
    return {"valid": True, "project_id": project_id}

# ─── Daily Reports (ElevenLabs) ───────────────────────────────────────────────
import base64

class DailyReportExplainRequest(BaseModel):
    project_id: str
    user_first_name: str
    date_str: str  # YYYY-MM-DD
    tasks: list
    force_regenerate: bool = False

@app.post("/api/reports/daily/explain")
async def explain_daily_report(req: DailyReportExplainRequest, authorization: str = Header(None)):
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Unauthorized")
    
    # 1. Check if report already exists for this date and project
    if not req.force_regenerate:
        try:
            report_resp = supabase.table("dmz_agents_reports").select("*").eq("project_id", req.project_id).eq("report_date", req.date_str).execute()
            if report_resp.data:
                existing_report = report_resp.data[0]
                # Since generating audio takes time and money, if we have it, return it.
                return {
                    "script": existing_report.get("narrative"), 
                    "audioUrl": existing_report.get("audio_url"),
                    "is_cached": True
                }
        except Exception as e:
            print(f"Error checking existing report: {e}")
            pass # Ignore and generate new if check failed

    # If no tasks and no report generated yet, return simple message
    if not req.tasks:
        return {"script": f"Oi {req.user_first_name.split()[0].title()}, parece que não tivemos nenhuma atividade registrada para este dia no projeto.", "audioUrl": None}

    first_name_only = req.user_first_name.split()[0].title()

    # 2. Format Date for Narration (PT-BR)
    try:
        from datetime import datetime
        dt = datetime.strptime(req.date_str, "%Y-%m-%d")
        weekdays = ["segunda-feira", "terça-feira", "quarta-feira", "quinta-feira", "sexta-feira", "sábado", "domingo"]
        months = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
        date_full = f"{weekdays[dt.weekday()]}, {dt.day} de {months[dt.month-1]} de {dt.year}"
    except:
        date_full = req.date_str

    # 3. Fetch Config from DB
    try:
        config_res = supabase.table("dmz_agents_config").select("value").eq("key", "reports_config").execute()
        if config_res.data:
            reports_config = config_res.data[0]["value"]
            system_prompt_template = reports_config.get("system_prompt", "")
            system_prompt = system_prompt_template.replace("{user_first_name}", first_name_only).replace("{date_str}", req.date_str).replace("{date_full}", date_full)
        else:
            raise Exception("Config not found")
    except Exception as e:
        print(f"Error fetching report config: {e}. Using fallback prompt.")
        system_prompt = (
            f"Você é a inteligência e analista (copywriter) do DMZ OS. O seu objetivo é construir uma narrativa em formato de relatório de "
            f"atividades. O gestor, {first_name_only}, pediu um resumo das tarefas do dia ({date_full}).\n\n"
            f"DIRETRIZES DE NARRATIVA:\n"
            f"1. Fale no plural em nome do 'squad DMZ'. Comece de forma amigável: 'Oi {first_name_only}, hoje o time trabalhou bastante...'.\n"
            f"2. PRIORIDADE: Primeiro destaque o que foi APROVADO (concluído de vez). Depois, dê ATENÇÃO ESPECIAL ao que está em 'done' (aguardando a aprovação dele). Em seguida, fale do que está em andamento (on_going), o que precisa de ajustes (rework) e o que está no radar para começar (to_do).\n"
            f"3. ESTILO: Crie uma história fluida, falando de forma natural. Não use bullets. Mencione os agentes pelo nome (orch, syd, oliver, etc) SEM usar o @.\n"
            f"4. FOCO EM RESULTADO: Explique o valor do que foi feito, não apenas o título da tarefa.\n"
            f"5. CONCLUSÃO: Encerre com uma frase motivadora ou um 'estamos à disposição'."
        )
    
    tasks_text = "; ".join([f"[{t.get('type')}] {t.get('title')} (responsável: {t.get('agent_handle', 'squad')})" for t in req.tasks])
    
    try:
        script = get_llm_response(system_prompt, f"Tarefas concluídas:\n{tasks_text}")
    except Exception as e:
        print(f"LLM Error generating narrative: {e}")
        raise HTTPException(status_code=500, detail=f"LLM Error: {str(e)}")
        
    # 3. Save to DB (Only the narrative)
    try:
        # If force_regenerate is True, we reset audio fields to trigger new audio generation in frontend
        report_data = {
            "project_id": req.project_id,
            "report_date": req.date_str,
            "narrative": script
        }
        
        if req.force_regenerate:
            report_data["has_audio"] = False
            report_data["audio_url"] = None

        supabase.table("dmz_agents_reports").upsert(
            report_data, 
            on_conflict="project_id,report_date"
        ).execute()
    except Exception as e:
        print(f"Failed to save report to DB: {e}")
        
    return {"script": script, "audioUrl": None, "is_cached": False}

# ─── API Key Management ─────────────────────────────────────────────────────

class KeyValidateRequest(BaseModel):
    api_key: str
    project_slug: str

class KeyGenerateRequest(BaseModel):
    project_id: str
    name: str

@app.post("/api/validate-key")
async def validate_key(req: KeyValidateRequest):
    # Hash incoming key
    key_hash = hashlib.sha256(req.api_key.encode()).hexdigest()
    
    # Check in DB
    try:
        res = supabase.table("dmz_agents_apikeys").select("*").eq("key_hash", key_hash).eq("project_id", req.project_slug).eq("is_active", True).execute()
        
        if not res.data:
            # Fallback: check if the key is the project_slug itself (legacy/dev support for easy migration)
            # Actually, better to strictly enforce keys now.
            raise HTTPException(status_code=401, detail="Invalid API Key or Project Slug association")
        
        key_data = res.data[0]
        
        # Update last used
        supabase.table("dmz_agents_apikeys").update({"last_used_at": datetime.utcnow().isoformat()}).eq("id", key_data["id"]).execute()
        
        return {"valid": True, "project_id": key_data["project_id"], "name": key_data["name"]}
    except Exception as e:
        print(f"Validation error: {e}")
        raise HTTPException(status_code=401, detail="Unauthorized")

@app.post("/api/keys/generate")
async def generate_key(req: KeyGenerateRequest):
    prefix = "dmz_pk_"
    raw_key = prefix + secrets.token_urlsafe(32)
    key_hash = hashlib.sha256(raw_key.encode()).hexdigest()
    
    try:
        new_key = {
            "project_id": req.project_id,
            "name": req.name,
            "key_hash": key_hash,
            "key_prefix": prefix,
            "is_active": True,
            "created_at": datetime.utcnow().isoformat()
        }
        
        res = supabase.table("dmz_agents_apikeys").insert(new_key).execute()
        
        if not res.data:
            raise HTTPException(status_code=500, detail="Failed to generate key")
            
        return {"api_key": raw_key, "id": res.data[0]["id"]}
    except Exception as e:
        print(f"Key generation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ─── MCP Backend Processor ───────────────────────────────────────────────────
# Processa tasks criadas via MCP Server (source: mcp_server) automaticamente

class MCPUpdateTaskRequest(BaseModel):
    task_id: str
    coluna: str | None = None
    status: str | None = None
    agent_id: str | None = None
    priority: int | None = None
    feedback: str | None = None

class MCPCommentRequest(BaseModel):
    task_id: str
    comment: str
    author: str = "mcp_server"

async def _process_mcp_task(task: dict):
    """Processa uma task MCP: carrega prompt do agente, envia ao LLM, grava response."""
    agent_db_id = task.get("agent_id") or "orchestrator"
    
    # Marca como in_progress
    supabase.table("dmz_agents_tasks").update({"status": "in_progress"}).eq("id", task["id"]).execute()
    
    # Carrega prompt do agente
    system_prompt = get_agent_prompt(agent_db_id)
    
    # Injeta contexto temporal
    from datetime import timezone, timedelta
    tz_sp = timezone(timedelta(hours=-3))
    now_sp = datetime.now(tz_sp)
    system_prompt += f"\n\nData atual: {now_sp.strftime('%d/%m/%Y %H:%M')} (GMT-3)"
    system_prompt += "\n\nVocê está respondendo uma solicitação feita via MCP Server (IDE do desenvolvedor). Seja direto, técnico e executivo na resposta."
    
    user_prompt = f"### DEMANDA\n\nTÍTULO: {task['title']}\nDESCRIÇÃO: {task.get('description', 'Sem descrição adicional.')}\n\nResponda de forma completa e técnica."
    
    try:
        response_text = get_llm_response(system_prompt, user_prompt)
        
        supabase.table("dmz_agents_tasks").update({
            "status": "completed",
            "completed_at": datetime.now(timezone.utc).isoformat(),
            "response": response_text,
            "responded_at": datetime.now(timezone.utc).isoformat(),
            "feedback": f"Processado via Backend (Railway) por @{agent_db_id}",
        }).eq("id", task["id"]).execute()
        
        print(f"[MCP] ✓ Task processada: {task['title'][:60]} → @{agent_db_id}")
        return True
    except Exception as e:
        supabase.table("dmz_agents_tasks").update({
            "status": "blocked",
            "response": f"Erro ao processar: {str(e)}",
            "responded_at": datetime.now(timezone.utc).isoformat(),
            "feedback": f"Erro Backend: {str(e)}",
        }).eq("id", task["id"]).execute()
        print(f"[MCP] ✖ Erro: {task['title'][:40]} — {e}")
        return False


@app.post("/mcp/process")
async def mcp_process_pending():
    """Processa tasks MCP pendentes (chamado pelo poller ou manualmente)."""
    try:
        # Busca tasks com source: mcp_server que ainda estão pendentes
        res = supabase.table("dmz_agents_tasks")\
            .select("*")\
            .eq("status", "pending")\
            .is_("response", "null")\
            .order("created_at")\
            .limit(5)\
            .execute()
        
        if not res.data:
            return {"processed": 0, "message": "Nenhuma task MCP pendente"}
        
        # Filtra só as que vieram do MCP
        mcp_tasks = [t for t in res.data if (t.get("metadata") or {}).get("source") == "mcp_server"]
        
        if not mcp_tasks:
            return {"processed": 0, "message": "Nenhuma task MCP pendente"}
        
        processed = 0
        for task in mcp_tasks:
            success = await _process_mcp_task(task)
            if success:
                processed += 1
        
        return {"processed": processed, "total_found": len(mcp_tasks)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/mcp/update-task")
async def mcp_update_task(req: MCPUpdateTaskRequest):
    """Atualiza uma task do Kanban (coluna, status, agente, prioridade)."""
    try:
        update_data = {}
        STATUS_MAP = {"master_plan": "pending", "to_do": "pending", "on_going": "in_progress", "done": "completed", "rework": "in_progress", "approved": "completed"}
        
        if req.coluna:
            update_data["type"] = req.coluna
            update_data["status"] = STATUS_MAP.get(req.coluna, "pending")
            if req.coluna in ("done", "approved"):
                update_data["completed_at"] = datetime.utcnow().isoformat()
            else:
                update_data["completed_at"] = None
        if req.status:
            update_data["status"] = req.status
        if req.agent_id is not None:
            update_data["agent_id"] = req.agent_id if req.agent_id else None
        if req.priority is not None:
            update_data["priority"] = req.priority
        if req.feedback is not None:
            update_data["feedback"] = req.feedback
        
        if not update_data:
            return {"error": "Nenhum campo para atualizar"}
        
        supabase.table("dmz_agents_tasks").update(update_data).eq("id", req.task_id).execute()
        return {"success": True, "task_id": req.task_id, "updated_fields": list(update_data.keys())}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/mcp/comment")
async def mcp_add_comment(req: MCPCommentRequest):
    """Adiciona um comentário ao feedback de uma task."""
    try:
        # Busca feedback atual
        res = supabase.table("dmz_agents_tasks").select("feedback").eq("id", req.task_id).single().execute()
        current = res.data.get("feedback") or ""
        
        # Append comment com timestamp
        from datetime import timezone, timedelta
        tz_sp = timezone(timedelta(hours=-3))
        now = datetime.now(tz_sp).strftime("%d/%m %H:%M")
        
        new_feedback = f"{current}\n[{now}] @{req.author}: {req.comment}".strip()
        
        supabase.table("dmz_agents_tasks").update({"feedback": new_feedback}).eq("id", req.task_id).execute()
        return {"success": True, "task_id": req.task_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── Telegram Webhook ───────────────────────────────────────────────────────
from fastapi import Request, BackgroundTasks
from webhook_telegram import handle_telegram_webhook

@app.post("/webhook/telegram")
async def telegram_webhook(req: Request, background_tasks: BackgroundTasks):
    """Webhook do Telegram para a Agente Yvi."""
    # Processamos em background para liberar o Telegram imediatamente
    background_tasks.add_task(handle_telegram_webhook, req, supabase, get_llm_response, get_sys_prompt)
    return {"status": "accepted"}

# ─── Background MCP Poller ──────────────────────────────────────────────────
import asyncio
import threading

def _mcp_poller_loop():
    """Background thread que checa tasks MCP pendentes a cada 10 segundos."""
    import time as _time
    import requests as _requests
    
    port = int(os.getenv("PORT", 8080))
    _time.sleep(5)  # Espera o server inicializar
    print("[MCP Poller] Iniciado — checando tasks MCP a cada 10s")
    
    while True:
        try:
            resp = _requests.post(f"http://localhost:{port}/mcp/process", timeout=120)
            data = resp.json()
            if data.get("processed", 0) > 0:
                print(f"[MCP Poller] Processou {data['processed']} task(s)")
        except Exception as e:
            pass  # Silencioso em caso de erro transitório
        _time.sleep(10)

# Inicia o poller em background ao startar o server
_poller_thread = threading.Thread(target=_mcp_poller_loop, daemon=True)
_poller_thread.start()


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", 8080)))
